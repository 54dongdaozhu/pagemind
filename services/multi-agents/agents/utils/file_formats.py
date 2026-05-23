import os
import re
import logging

logger = logging.getLogger(__name__)


def markdown_to_docx(markdown_text: str, output_path: str) -> None:
    """Convert a Markdown string to a .docx file at output_path."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Basic style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in markdown_text.splitlines():
        stripped = line.rstrip()

        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 40)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Pt(24)
            p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif stripped == "":
            doc.add_paragraph("")
        else:
            _add_inline_paragraph(doc, stripped)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    logger.info("Saved docx to %s", output_path)


def _add_inline_paragraph(doc, text: str):
    """Add a paragraph handling **bold** and *italic* inline markers."""
    from docx import Document

    p = doc.add_paragraph()
    # Split on bold/italic markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            p.add_run(part)
    return p
